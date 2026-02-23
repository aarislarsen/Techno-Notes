from flask import Flask, request, jsonify, render_template, send_file
from werkzeug.utils import secure_filename
import requests
import os
import json
import subprocess
import sys
import time
import threading
from pathlib import Path
import platform
import shutil
import secrets
import signal
import logging
from datetime import datetime
from functools import wraps
import re
import zipfile

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['OUTPUT_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'outputs')
app.config['SECRET_KEY'] = secrets.token_hex(32)
PROMPT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'prompt.txt')
CONFIG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'llm_config.json')
API_KEYS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.api_keys.json')
TEMPLATE_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'word_templates')
TEMPLATE_FILE = os.path.join(TEMPLATE_FOLDER, 'template.docx')
TIMING_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'processing_times.json')
OLLAMA_URL = 'http://127.0.0.1:11434'
ALLOWED_EXTENSIONS = {'pdf', 'zip'}

# Rate limiting - simple in-memory store
request_times = {}
RATE_LIMIT = 10  # requests per minute
RATE_WINDOW = 60  # seconds

# Create directories with proper permissions
for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER'], 'logs', TEMPLATE_FOLDER]:
    os.makedirs(folder, mode=0o700, exist_ok=True)

# Whitelist of allowed models
ALLOWED_MODELS = {
    'llama2:13b', 'llama2:70b',
    'llama3', 'llama3:8b', 'llama3.2:1b', 'llama3.2:3b',
    'llama3.2-vision', 'llama3.2-vision:11b',
    'mistral-nemo', 'mistral-nemo:12b',
    'qwen2.5', 'qwen2.5:7b', 'qwen2.5:14b',
    'command-r', 'command-r:35b',
    'phi', 'phi:2.7b', 'phi3',
    'codellama', 'codellama:7b',
    'gemma', 'gemma:2b', 'gemma:7b'
}

def rate_limit(f):
    """Simple rate limiting decorator"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        client_ip = request.remote_addr
        current_time = time.time()
        
        # Clean old entries
        if client_ip in request_times:
            request_times[client_ip] = [t for t in request_times[client_ip] 
                                       if current_time - t < RATE_WINDOW]
        
        # Check rate limit
        if client_ip in request_times and len(request_times[client_ip]) >= RATE_LIMIT:
            logger.warning(f"Rate limit exceeded for {client_ip}")
            return jsonify({'error': 'Rate limit exceeded. Try again later.'}), 429
        
        # Add current request
        if client_ip not in request_times:
            request_times[client_ip] = []
        request_times[client_ip].append(current_time)
        
        return f(*args, **kwargs)
    return decorated_function

def validate_filename(filename):
    """Validate filename to prevent path traversal"""
    if not filename or not isinstance(filename, str):
        return False
    filename = os.path.basename(filename)
    if '..' in filename or '/' in filename or '\\' in filename:
        return False
    if len(filename) > 255:
        return False
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def validate_model_name(model_name):
    """Validate model name against whitelist"""
    if not model_name or not isinstance(model_name, str):
        return False
    model_name = model_name.lower().strip()
    if len(model_name) > 50:
        return False
    for allowed in ALLOWED_MODELS:
        if model_name == allowed or model_name.startswith(allowed + ':'):
            return True
    return False

def sanitize_input(text, max_length=10000):
    """Sanitize text input for HTML contexts"""
    if not isinstance(text, str):
        return ""
    text = text[:max_length]
    text = text.replace('\x00', '')
    # Remove potential XSS vectors
    text = text.replace('<', '&lt;').replace('>', '&gt;')
    return text


def sanitize_text(text, max_length=10000):
    """Sanitize plain text (no HTML escaping) for non-HTML contexts like LLM prompts"""
    if not isinstance(text, str):
        return ""
    text = text[:max_length]
    text = text.replace('\x00', '')
    return text

def safe_remove_file(filepath):
    """Safely remove a file with error handling"""
    try:
        if os.path.exists(filepath):
            os.remove(filepath)
            return True
    except OSError as e:
        logger.error(f"Failed to remove file {filepath}: {e}")
    return False

class LLMManager:
    def __init__(self):
        self.config = self.load_config()
        self.ollama_process = None
        self.setup_lock = threading.Lock()
        self.setup_progress = {
            'stage': 'idle',
            'message': '',
            'progress': 0,
            'error': None
        }
        self.model_download_progress = {
            'active': False,
            'model': '',
            'status': '',
            'total': 0,
            'completed': 0,
            'percent': 0,
            'started_at': None,
            'eta_seconds': None,
            'error': None
        }
    
    def load_config(self):
        """Safely load configuration"""
        try:
            if os.path.exists(CONFIG_FILE):
                with open(CONFIG_FILE, 'r') as f:
                    config = json.load(f)
                    if not isinstance(config, dict):
                        raise ValueError("Invalid config format")
                    if 'model_name' in config and not validate_model_name(config['model_name']):
                        config['model_name'] = 'qwen2.5:14b'
                    return config
        except (json.JSONDecodeError, ValueError, IOError) as e:
            logger.error(f"Config load error: {e}")
        
        return {
            'ollama_installed': False,
            'ollama_path': None,
            'model_name': 'qwen2.5:14b',
            'model_downloaded': False,
            'setup_complete': False,
            'provider': 'ollama',  # Default to local ollama
            'api_model': 'gpt-4o-mini'  # Default API model
        }
    
    def save_config(self):
        """Safely save configuration"""
        try:
            with open(CONFIG_FILE, 'w') as f:
                json.dump(self.config, f, indent=2)
            try:
                os.chmod(CONFIG_FILE, 0o600)
            except OSError:
                pass  # Ignore on Windows/WSL
        except IOError as e:
            logger.error(f"Config save error: {e}")
    
    def update_progress(self, stage, message, progress):
        """Update setup progress with sanitized values"""
        self.setup_progress = {
            'stage': sanitize_input(str(stage), 50),
            'message': sanitize_input(str(message), 200),
            'progress': max(0, min(100, int(progress))),
            'error': None
        }
        logger.info(f"Setup progress: {stage} - {message} ({progress}%)")
    
    def check_ollama_installed(self):
        """Check if Ollama is installed"""
        try:
            result = subprocess.run(
                ['ollama', '--version'], 
                capture_output=True, 
                text=True, 
                timeout=5,
                check=False
            )
            if result.returncode == 0:
                self.config['ollama_installed'] = True
                self.config['ollama_path'] = self.get_ollama_path()
                self.save_config()
                logger.info(f"Ollama found: {result.stdout.strip()}")
                return True
        except (FileNotFoundError, subprocess.TimeoutExpired, OSError) as e:
            logger.debug(f"Ollama not found: {e}")
        return False
    
    def get_ollama_path(self):
        """Get Ollama executable path"""
        try:
            result = subprocess.run(
                ['which', 'ollama'],
                capture_output=True, 
                text=True,
                timeout=5,
                check=False
            )
            if result.returncode == 0:
                path = result.stdout.strip().split('\n')[0]
                if os.path.isfile(path):
                    return path
        except (subprocess.TimeoutExpired, OSError):
            pass
        return None
    
    def install_ollama(self):
        """Automatically install Ollama"""
        self.update_progress('installing', 'Installing Ollama...', 20)
        
        try:
            install_url = 'https://ollama.com/install.sh'
            install_script = '/tmp/ollama_install.sh'
            
            # Download install script
            self.update_progress('installing', 'Downloading Ollama installer...', 25)
            logger.info("Downloading Ollama install script")
            
            result = subprocess.run([
                'curl', '-fsSL', '--max-time', '60', '--max-filesize', '1000000',
                install_url, '-o', install_script
            ], check=True, timeout=65, capture_output=True)
            
            if not os.path.exists(install_script):
                raise Exception("Install script download failed")
            
            file_size = os.path.getsize(install_script)
            if file_size > 1000000 or file_size < 100:
                safe_remove_file(install_script)
                raise Exception(f"Install script size suspicious: {file_size} bytes")
            
            try:
                os.chmod(install_script, 0o700)
            except OSError:
                pass
            
            # Run install script
            self.update_progress('installing', 'Running Ollama installer...', 40)
            logger.info("Running Ollama installer")

            result = subprocess.run(
                ['sh', install_script],
                check=True,
                timeout=600,
                env={'PATH': os.environ.get('PATH', '/usr/local/bin:/usr/bin:/bin')},
                stdin=subprocess.DEVNULL,
                capture_output=True,
                text=True
            )
            
            logger.info(f"Ollama installer output: {result.stdout[:200]}")
            
            safe_remove_file(install_script)
            
            time.sleep(2)
            if self.check_ollama_installed():
                self.update_progress('installed', 'Ollama installed successfully', 50)
                return True
            else:
                raise Exception("Installation completed but ollama command not found in PATH")
                
        except subprocess.TimeoutExpired:
            self.setup_progress['error'] = "Installation timeout - check network connection"
            logger.error("Ollama installation timeout")
            return False
        except subprocess.CalledProcessError as e:
            error_msg = e.stderr if hasattr(e, 'stderr') and e.stderr else str(e)
            self.setup_progress['error'] = f"Installation failed: {error_msg[:100]}"
            logger.error(f"Ollama installation failed: {error_msg}")
            return False
        except Exception as e:
            self.setup_progress['error'] = str(e)[:200]
            logger.error(f"Ollama installation error: {e}")
            return False
    
    def start_ollama_service(self):
        """Start Ollama service"""
        if self.ollama_process and self.ollama_process.poll() is None:
            logger.info("Ollama service already running")
            return True
        
        try:
            self.update_progress('starting', 'Starting Ollama service...', 60)
            logger.info("Starting Ollama service")
            
            # Kill existing processes
            try:
                subprocess.run(
                    ['pkill', '-9', 'ollama'], 
                    capture_output=True,
                    timeout=5,
                    check=False
                )
                time.sleep(2)
            except (subprocess.TimeoutExpired, OSError):
                pass
            
            # Start service
            env = os.environ.copy()
            env['OLLAMA_HOST'] = '127.0.0.1:11434'
            
            self.ollama_process = subprocess.Popen(
                ['ollama', 'serve'],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                env=env,
                start_new_session=True
            )
            
            logger.info(f"Ollama process started with PID: {self.ollama_process.pid}")
            
            # Wait for service with better error handling
            for i in range(30):
                try:
                    response = requests.get(f'{OLLAMA_URL}/api/tags', timeout=2)
                    if response.status_code == 200:
                        self.update_progress('started', 'Ollama service running', 70)
                        logger.info("Ollama service is ready")
                        return True
                except requests.RequestException:
                    time.sleep(1)
            
            # Check if process died
            if self.ollama_process.poll() is not None:
                stderr = self.ollama_process.stderr.read().decode('utf-8')[:500]
                logger.error(f"Ollama failed to start: {stderr}")
                raise Exception(f"Ollama service failed to start: {stderr[:200]}")
            
            logger.error("Ollama service didn't respond in 30 seconds")
            return False
            
        except Exception as e:
            self.setup_progress['error'] = str(e)[:200]
            logger.error(f"Failed to start Ollama service: {e}")
            return False
    
    def check_model_downloaded(self, model_name=None):
        """Check if model is downloaded"""
        if model_name is None:
            model_name = self.config['model_name']
        
        if not validate_model_name(model_name):
            return False
        
        try:
            response = requests.get(f'{OLLAMA_URL}/api/tags', timeout=5)
            if response.status_code == 200:
                data = response.json()
                models = data.get('models', [])
                for model in models:
                    if model_name in model.get('name', ''):
                        self.config['model_downloaded'] = True
                        self.save_config()
                        logger.info(f"Model {model_name} found")
                        return True
        except (requests.RequestException, json.JSONDecodeError) as e:
            logger.debug(f"Error checking model: {e}")
        return False
    
    def download_model(self, model_name=None):
        """Download model with validation"""
        if model_name is None:
            model_name = self.config['model_name']
        
        if not validate_model_name(model_name):
            self.setup_progress['error'] = "Invalid model name"
            return False
        
        self.update_progress('downloading', f'Downloading model {model_name}...', 80)
        logger.info(f"Starting download of model: {model_name}")
        
        try:
            process = subprocess.Popen(
                ['ollama', 'pull', model_name],
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                bufsize=1
            )
            
            start_time = time.time()
            last_update = time.time()
            
            for line in process.stdout:
                current_time = time.time()
                
                # Timeout check (30 minutes)
                if current_time - start_time > 1800:
                    process.kill()
                    raise Exception("Download timeout (30 minutes exceeded)")
                
                # Update progress message every 2 seconds
                if current_time - last_update > 2:
                    if 'pulling' in line.lower() or 'downloading' in line.lower():
                        clean_line = line.strip()[:200]
                        self.setup_progress['message'] = sanitize_input(clean_line, 200)
                        logger.debug(f"Download progress: {clean_line}")
                        last_update = current_time
            
            process.wait(timeout=60)
            
            if process.returncode == 0:
                self.config['model_name'] = model_name
                self.config['model_downloaded'] = True
                self.save_config()
                self.update_progress('complete', f'Model {model_name} ready', 95)
                logger.info(f"Model {model_name} downloaded successfully")
                return True
            else:
                raise Exception(f"Download failed with exit code {process.returncode}")
                
        except subprocess.TimeoutExpired:
            self.setup_progress['error'] = "Download timeout"
            logger.error("Model download timeout")
            return False
        except Exception as e:
            self.setup_progress['error'] = str(e)[:200]
            logger.error(f"Model download error: {e}")
            return False
    
    def check_model_available(self, model_name):
        """Check if a model is already available locally via Ollama API"""
        try:
            response = requests.get(f'{OLLAMA_URL}/api/tags', timeout=5)
            if response.status_code == 200:
                models = response.json().get('models', [])
                for m in models:
                    name = m.get('name', '')
                    # Match "llama2:latest" against "llama2", or exact match
                    if name == model_name or name.startswith(model_name + ':') or name == model_name + ':latest':
                        return True
        except requests.RequestException:
            pass
        return False

    def pull_model_async(self, model_name):
        """Download a model via Ollama REST API with streaming progress"""
        self.model_download_progress = {
            'active': True,
            'model': model_name,
            'status': 'downloading',
            'total': 0,
            'completed': 0,
            'percent': 0,
            'started_at': time.time(),
            'eta_seconds': None,
            'error': None
        }

        try:
            response = requests.post(
                f'{OLLAMA_URL}/api/pull',
                json={'name': model_name, 'stream': True},
                stream=True,
                timeout=3600
            )
            response.raise_for_status()

            for line in response.iter_lines():
                if not line:
                    continue
                try:
                    data = json.loads(line)
                except json.JSONDecodeError:
                    continue

                status = data.get('status', '')
                total = data.get('total', 0)
                completed = data.get('completed', 0)

                if total and total > 0:
                    self.model_download_progress['total'] = total
                    self.model_download_progress['completed'] = completed
                    self.model_download_progress['percent'] = min(99, int(completed * 100 / total))

                    elapsed = time.time() - self.model_download_progress['started_at']
                    if completed > 0 and elapsed > 2:
                        remaining_bytes = total - completed
                        speed = completed / elapsed
                        if speed > 0:
                            self.model_download_progress['eta_seconds'] = round(remaining_bytes / speed)

                self.model_download_progress['status_message'] = status

                if data.get('error'):
                    raise Exception(data['error'])

            # Download complete
            self.model_download_progress['percent'] = 100
            self.model_download_progress['eta_seconds'] = 0
            self.model_download_progress['status'] = 'complete'
            self.model_download_progress['active'] = False

            self.config['model_name'] = model_name
            self.config['model_downloaded'] = True
            self.save_config()
            logger.info(f"Model {model_name} pulled successfully")

        except Exception as e:
            logger.error(f"Model pull error: {e}")
            self.model_download_progress['status'] = 'error'
            self.model_download_progress['error'] = str(e)[:200]
            self.model_download_progress['active'] = False

    def get_status(self):
        """Get system status"""
        ollama_running = False
        try:
            response = requests.get(f'{OLLAMA_URL}/api/tags', timeout=2)
            ollama_running = response.status_code == 200
        except requests.RequestException:
            pass
        
        return {
            'ollama_installed': self.config['ollama_installed'],
            'ollama_running': ollama_running,
            'model_downloaded': self.config['model_downloaded'],
            'model_name': self.config['model_name'],
            'setup_complete': self.config.get('setup_complete', False),
            'ready': self.config['ollama_installed'] and ollama_running and self.config['model_downloaded']
        }
    
    def auto_setup(self, model_name='qwen2.5:14b'):
        """Automated model download and service startup"""
        if not validate_model_name(model_name):
            self.setup_progress['error'] = "Invalid model name"
            logger.error(f"Invalid model name requested: {model_name}")
            return False

        with self.setup_lock:
            try:
                logger.info(f"Starting auto-setup for model: {model_name}")

                # Step 1: Check Ollama is installed
                if not self.check_ollama_installed():
                    self.update_progress('error', 'Ollama not installed. Please run ./setup.sh first', 0)
                    logger.error("Ollama not installed - setup.sh must be run first")
                    return False

                self.update_progress('checking', 'Ollama detected', 20)

                # Step 2: Start service
                if not self.start_ollama_service():
                    logger.error("Failed to start Ollama service")
                    return False

                # Step 3: Download model
                if not self.check_model_downloaded(model_name):
                    if not self.download_model(model_name):
                        logger.error("Model download failed")
                        return False
                else:
                    self.update_progress('checking', f'Model {model_name} already downloaded', 90)

                # Mark complete
                self.config['setup_complete'] = True
                self.save_config()
                self.update_progress('complete', 'Setup complete - System ready', 100)
                logger.info("Auto-setup completed successfully")
                return True

            except Exception as e:
                self.setup_progress['error'] = str(e)[:200]
                logger.error(f"Auto-setup error: {e}")
                return False
    
    def cleanup(self):
        """Cleanup on shutdown"""
        if self.ollama_process and self.ollama_process.poll() is None:
            try:
                logger.info("Terminating Ollama process")
                self.ollama_process.terminate()
                self.ollama_process.wait(timeout=5)
            except:
                try:
                    self.ollama_process.kill()
                except:
                    pass

llm_manager = LLMManager()

def read_prompt():
    """Safely read prompt file"""
    try:
        if os.path.exists(PROMPT_FILE):
            with open(PROMPT_FILE, 'r', encoding='utf-8') as f:
                content = f.read(50000)
                return sanitize_text(content, 10000)
    except IOError as e:
        logger.error(f"Failed to read prompt file: {e}")
    return "Analyze this PDF document and provide a summary."

def save_prompt(prompt_text):
    """Safely save prompt file"""
    try:
        prompt_text = sanitize_text(prompt_text, 10000)
        with open(PROMPT_FILE, 'w', encoding='utf-8') as f:
            f.write(prompt_text)
        try:
            os.chmod(PROMPT_FILE, 0o600)
        except OSError:
            pass
        logger.info("Prompt saved successfully")
        return True
    except IOError as e:
        logger.error(f"Failed to save prompt: {e}")
        raise Exception(f"Failed to save prompt: {e}")

def load_processing_times():
    """Load processing time history from disk"""
    try:
        if os.path.exists(TIMING_FILE):
            with open(TIMING_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if isinstance(data, list):
                    return [float(t) for t in data if isinstance(t, (int, float))]
    except (IOError, json.JSONDecodeError, ValueError) as e:
        logger.error(f"Failed to load processing times: {e}")
    return []


def save_processing_time(duration_seconds):
    """Append a processing duration and save (keep last 50)"""
    try:
        times = load_processing_times()
        times.append(round(duration_seconds, 2))
        times = times[-50:]
        with open(TIMING_FILE, 'w', encoding='utf-8') as f:
            json.dump(times, f)
    except IOError as e:
        logger.error(f"Failed to save processing time: {e}")


def get_average_processing_time():
    """Return the average processing time in seconds, or None if no history"""
    times = load_processing_times()
    if not times:
        return None
    return round(sum(times) / len(times), 1)


def extract_template_structure(template_path):
    """Extract the structure of a Word template for LLM guidance"""
    try:
        from docx import Document
        doc = Document(template_path)
        structure_lines = []
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else 'Normal'
            text = para.text.strip()
            if text:
                structure_lines.append(f"[{style_name}] {text}")
            elif style_name != 'Normal':
                structure_lines.append(f"[{style_name}] (empty)")
        if not structure_lines:
            return ""
        return (
            "IMPORTANT: Format your response following this document template structure. "
            "Use markdown headings (#, ##, ###) for headings, bullet points (- ) for lists, "
            "and plain text for normal paragraphs. Match the section order and hierarchy below:\n\n"
            + "\n".join(structure_lines)
        )
    except Exception as e:
        logger.error(f"Failed to extract template structure: {e}")
        return ""


def generate_docx_from_template(template_path, llm_response):
    """Generate a .docx file from the template with LLM response content"""
    from docx import Document
    from docx.shared import Pt

    doc = Document(template_path)

    # Clear existing template content
    for para in doc.paragraphs:
        p_element = para._element
        p_element.getparent().remove(p_element)

    # Also clear any tables from template
    for table in doc.tables:
        t_element = table._element
        t_element.getparent().remove(t_element)

    # Collect available style names for safe fallback
    available_styles = {s.name for s in doc.styles}

    def get_style(preferred, fallback='Normal'):
        return preferred if preferred in available_styles else fallback

    # Parse LLM response (markdown-like) and add to document
    lines = llm_response.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Heading 1: # Heading
        if stripped.startswith('# ') and not stripped.startswith('## '):
            heading_text = stripped[2:].strip()
            heading_text = re.sub(r'\*\*(.+?)\*\*', r'\1', heading_text)
            doc.add_paragraph(heading_text, style=get_style('Heading 1'))

        # Heading 2: ## Heading
        elif stripped.startswith('## ') and not stripped.startswith('### '):
            heading_text = stripped[3:].strip()
            heading_text = re.sub(r'\*\*(.+?)\*\*', r'\1', heading_text)
            doc.add_paragraph(heading_text, style=get_style('Heading 2'))

        # Heading 3: ### Heading
        elif stripped.startswith('### '):
            heading_text = stripped[4:].strip()
            heading_text = re.sub(r'\*\*(.+?)\*\*', r'\1', heading_text)
            doc.add_paragraph(heading_text, style=get_style('Heading 3'))

        # Bullet list: - item or * item
        elif re.match(r'^[-*]\s+', stripped):
            item_text = re.sub(r'^[-*]\s+', '', stripped)
            para = doc.add_paragraph(style=get_style('List Bullet'))
            _add_formatted_runs(para, item_text)

        # Numbered list: 1. item
        elif re.match(r'^\d+\.\s+', stripped):
            item_text = re.sub(r'^\d+\.\s+', '', stripped)
            para = doc.add_paragraph(style=get_style('List Number'))
            _add_formatted_runs(para, item_text)

        # Normal paragraph
        else:
            para = doc.add_paragraph(style=get_style('Normal'))
            _add_formatted_runs(para, stripped)

        i += 1

    return doc


def _add_formatted_runs(paragraph, text):
    """Add text to a paragraph, handling **bold** and *italic* markdown"""
    # Split on bold (**text**) and italic (*text*) patterns
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def extract_pdf_text(pdf_path):
    """Safely extract PDF text"""
    try:
        import PyPDF2
    except ImportError:
        raise Exception("PyPDF2 not installed")
    
    if not os.path.isfile(pdf_path):
        raise Exception("PDF file not found")
    
    file_size = os.path.getsize(pdf_path)
    if file_size > 50 * 1024 * 1024:
        raise Exception("PDF file too large (max 50MB)")
    
    logger.info(f"Extracting text from PDF ({file_size} bytes)")
    text = ""
    
    try:
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            num_pages = len(reader.pages)
            max_pages = min(num_pages, 100)
            
            logger.info(f"PDF has {num_pages} pages, processing {max_pages}")
            
            for i in range(max_pages):
                try:
                    page_text = reader.pages[i].extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Failed to extract page {i+1}: {e}")
                    continue
                
                if len(text) > 500000:
                    logger.info("Text limit reached (500KB)")
                    text = text[:500000]
                    break
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise Exception(f"PDF extraction failed: {str(e)[:100]}")
    
    if not text.strip():
        raise Exception("No text could be extracted from PDF")
    
    logger.info(f"Extracted {len(text)} characters from PDF")
    return sanitize_input(text, 500000)

def extract_latex_text(zip_path):
    """Extract readable text from LaTeX .tex files inside a ZIP archive"""
    if not os.path.isfile(zip_path):
        raise Exception("ZIP file not found")

    if not zipfile.is_zipfile(zip_path):
        raise Exception("Invalid ZIP file")

    logger.info(f"Extracting LaTeX text from ZIP ({os.path.getsize(zip_path)} bytes)")
    texts = []
    total_size = 0
    max_uncompressed = 100 * 1024 * 1024  # 100MB uncompressed limit

    try:
        with zipfile.ZipFile(zip_path, 'r') as zf:
            # Check for zip bomb
            for info in zf.infolist():
                total_size += info.file_size
                if total_size > max_uncompressed:
                    raise Exception("ZIP contents too large (over 100MB uncompressed)")

            # Find and read .tex files
            tex_files = [f for f in zf.namelist()
                         if f.lower().endswith('.tex')
                         and '..' not in f
                         and not f.startswith('/')]

            if not tex_files:
                raise Exception("No .tex files found in ZIP archive")

            tex_files.sort()
            logger.info(f"Found {len(tex_files)} .tex files in ZIP")

            for tex_file in tex_files:
                try:
                    raw = zf.read(tex_file)
                    try:
                        content = raw.decode('utf-8')
                    except UnicodeDecodeError:
                        content = raw.decode('latin-1')

                    cleaned = _clean_latex(content)
                    if cleaned.strip():
                        texts.append(f"--- {tex_file} ---\n{cleaned}")
                except Exception as e:
                    logger.warning(f"Failed to read {tex_file}: {e}")
                    continue

                if sum(len(t) for t in texts) > 500000:
                    logger.info("Text limit reached (500KB)")
                    break

    except zipfile.BadZipFile:
        raise Exception("Corrupted ZIP file")
    except Exception as e:
        if "No .tex files" in str(e) or "too large" in str(e):
            raise
        logger.error(f"ZIP extraction failed: {e}")
        raise Exception(f"ZIP extraction failed: {str(e)[:100]}")

    combined = "\n\n".join(texts)
    if not combined.strip():
        raise Exception("No readable text found in LaTeX files")

    logger.info(f"Extracted {len(combined)} characters from LaTeX ZIP")
    return sanitize_text(combined, 500000)


def _clean_latex(text):
    """Convert LaTeX source to readable text, preserving structure"""
    # Remove comments
    text = re.sub(r'(?<!\\)%.*$', '', text, flags=re.MULTILINE)

    # Remove preamble (everything before \begin{document})
    doc_match = re.search(r'\\begin\{document\}', text)
    if doc_match:
        text = text[doc_match.end():]

    # Remove \end{document}
    text = re.sub(r'\\end\{document\}', '', text)

    # Convert structural commands to readable markers
    text = re.sub(r'\\chapter\*?\{([^}]*)\}', r'\n# \1\n', text)
    text = re.sub(r'\\section\*?\{([^}]*)\}', r'\n## \1\n', text)
    text = re.sub(r'\\subsection\*?\{([^}]*)\}', r'\n### \1\n', text)
    text = re.sub(r'\\subsubsection\*?\{([^}]*)\}', r'\n#### \1\n', text)
    text = re.sub(r'\\paragraph\*?\{([^}]*)\}', r'\n**\1** ', text)

    # Convert formatting to readable text
    text = re.sub(r'\\textbf\{([^}]*)\}', r'**\1**', text)
    text = re.sub(r'\\textit\{([^}]*)\}', r'*\1*', text)
    text = re.sub(r'\\emph\{([^}]*)\}', r'*\1*', text)
    text = re.sub(r'\\underline\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]*)\}', r'\1', text)

    # Convert list items
    text = re.sub(r'\\item\s*\[([^\]]*)\]', r'- \1: ', text)
    text = re.sub(r'\\item\s*', r'- ', text)

    # Remove environment wrappers but keep content
    text = re.sub(r'\\begin\{(itemize|enumerate|description|quote|quotation|center|flushleft|flushright|abstract|figure|table)\*?\}', '', text)
    text = re.sub(r'\\end\{(itemize|enumerate|description|quote|quotation|center|flushleft|flushright|abstract|figure|table)\*?\}', '', text)

    # Remove commands we don't need
    text = re.sub(r'\\(usepackage|documentclass|input|include|bibliography|bibliographystyle|label|ref|cite|pageref|eqref|url|href)\{[^}]*\}(\{[^}]*\})?', '', text)
    text = re.sub(r'\\(newcommand|renewcommand|def|let|setlength|setcounter|addtocounter)\{[^}]*\}.*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\\(vspace|hspace|vskip|hskip|bigskip|medskip|smallskip|noindent|clearpage|newpage|pagebreak|linebreak)\b\*?(\{[^}]*\})?', '', text)
    text = re.sub(r'\\(maketitle|tableofcontents|listoffigures|listoftables)\b', '', text)

    # Remove math environments but keep inline math readable
    text = re.sub(r'\$\$.*?\$\$', '[equation]', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{(equation|align|gather|multline)\*?\}.*?\\end\{\1\*?\}', '[equation]', text, flags=re.DOTALL)
    text = re.sub(r'\$([^$]+)\$', r'\1', text)

    # Remove remaining unknown commands (keep their arguments if present)
    text = re.sub(r'\\[a-zA-Z]+\*?\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+\*?', '', text)

    # Clean up braces and special chars
    text = text.replace('{', '').replace('}', '')
    text = text.replace('~', ' ')
    text = text.replace('\\\\', '\n')
    text = text.replace('\\&', '&')
    text = text.replace('\\%', '%')
    text = text.replace('\\$', '$')
    text = text.replace('\\_', '_')
    text = text.replace('\\#', '#')

    # Collapse excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)

    return text.strip()


def load_api_keys():
    """Load API keys from file"""
    try:
        if os.path.exists(API_KEYS_FILE):
            with open(API_KEYS_FILE, 'r') as f:
                keys = json.load(f)
                if isinstance(keys, dict):
                    return keys
    except (json.JSONDecodeError, IOError) as e:
        logger.error(f"API keys load error: {e}")
    return {}

def save_api_key(provider, api_key):
    """Save API key for a provider"""
    try:
        api_keys = load_api_keys()
        api_keys[provider] = api_key

        with open(API_KEYS_FILE, 'w') as f:
            json.dump(api_keys, f, indent=2)

        try:
            os.chmod(API_KEYS_FILE, 0o600)
        except OSError:
            pass

        logger.info(f"API key saved for provider: {provider}")
        return True
    except IOError as e:
        logger.error(f"API key save error: {e}")
        return False

def get_api_key(provider):
    """Get API key for a provider"""
    api_keys = load_api_keys()
    return api_keys.get(provider)

def query_chatgpt(prompt, context, model='gpt-4o-mini'):
    """Query ChatGPT API with timeout and validation"""
    prompt = sanitize_input(prompt, 10000)
    context = sanitize_input(context, 100000)

    api_key = get_api_key('chatgpt')
    if not api_key:
        raise Exception("ChatGPT API key not configured. Please add your API key in settings.")

    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json'
    }

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": prompt},
            {"role": "user", "content": f"Document content:\n\n{context}"}
        ],
        "temperature": 0.7
    }

    logger.info(f"Querying ChatGPT with model: {model}")

    try:
        response = requests.post(
            'https://api.openai.com/v1/chat/completions',
            headers=headers,
            json=payload,
            timeout=300
        )
        response.raise_for_status()

        result = response.json()
        llm_response = result['choices'][0]['message']['content']

        logger.info(f"Received response from ChatGPT ({len(llm_response)} characters)")
        return sanitize_input(llm_response, 1000000)

    except requests.Timeout:
        logger.error("ChatGPT request timeout")
        raise Exception("ChatGPT request timeout (5 minutes)")
    except requests.RequestException as e:
        logger.error(f"ChatGPT request failed: {e}")
        error_detail = ""
        if hasattr(e, 'response') and e.response is not None:
            try:
                error_detail = e.response.json().get('error', {}).get('message', '')
            except:
                pass
        raise Exception(f"ChatGPT API error: {error_detail or str(e)[:100]}")
    except (json.JSONDecodeError, KeyError) as e:
        logger.error(f"Invalid ChatGPT response: {e}")
        raise Exception("Invalid ChatGPT response format")

def query_perplexity(prompt, context, model='llama-3.1-sonar-small-128k-online'):
    """Query Perplexity API with timeout and validation"""
    prompt = sanitize_input(prompt, 10000)
    context = sanitize_input(context, 100000)

    api_key = get_api_key('perplexity')
    if not api_key:
        raise Exception("Perplexity API key not configured. Please add your API key in settings.")

    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json'
    }

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": prompt},
            {"role": "user", "content": f"Document content:\n\n{context}"}
        ]
    }

    logger.info(f"Querying Perplexity with model: {model}")

    try:
        response = requests.post(
            'https://api.perplexity.ai/chat/completions',
            headers=headers,
            json=payload,
            timeout=300
        )
        response.raise_for_status()

        result = response.json()
        llm_response = result['choices'][0]['message']['content']

        logger.info(f"Received response from Perplexity ({len(llm_response)} characters)")
        return sanitize_input(llm_response, 1000000)

    except requests.Timeout:
        logger.error("Perplexity request timeout")
        raise Exception("Perplexity request timeout (5 minutes)")
    except requests.RequestException as e:
        logger.error(f"Perplexity request failed: {e}")
        error_detail = ""
        if hasattr(e, 'response') and e.response is not None:
            try:
                error_detail = e.response.text[:200]
            except:
                pass
        raise Exception(f"Perplexity API error: {error_detail or str(e)[:100]}")
    except (json.JSONDecodeError, KeyError) as e:
        logger.error(f"Invalid Perplexity response: {e}")
        raise Exception("Invalid Perplexity response format")

def query_ollama(prompt, context):
    """Query Ollama with timeout and validation"""
    prompt = sanitize_input(prompt, 10000)
    context = sanitize_input(context, 500000)

    payload = {
        "model": llm_manager.config['model_name'],
        "prompt": f"{prompt}\n\nDocument content:\n{context}",
        "stream": False,
        "options": {
            "num_ctx": 131072  # 128K token context window — handles up to ~100 pages
        }
    }

    logger.info(f"Querying Ollama with model: {llm_manager.config['model_name']}")

    try:
        response = requests.post(
            f'{OLLAMA_URL}/api/generate',
            json=payload,
            timeout=1200
        )
        response.raise_for_status()

        result = response.json()
        llm_response = result.get('response', '')

        logger.info(f"Received response from Ollama ({len(llm_response)} characters)")
        return sanitize_input(llm_response, 1000000)

    except requests.Timeout:
        logger.error("Ollama request timeout")
        raise Exception("LLM request timeout (20 minutes)")
    except requests.RequestException as e:
        logger.error(f"Ollama request failed: {e}")
        raise Exception(f"LLM request failed: {str(e)[:100]}")
    except json.JSONDecodeError as e:
        logger.error(f"Invalid Ollama response: {e}")
        raise Exception("Invalid LLM response format")

def query_llm(prompt, context):
    """Route LLM query to appropriate provider"""
    provider = llm_manager.config.get('provider', 'ollama')
    api_model = llm_manager.config.get('api_model', 'gpt-4o-mini')

    logger.info(f"Routing query to provider: {provider}")

    if provider == 'chatgpt':
        return query_chatgpt(prompt, context, api_model)
    elif provider == 'perplexity':
        return query_perplexity(prompt, context, api_model)
    else:  # Default to ollama
        return query_ollama(prompt, context)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/setup_status', methods=['GET'])
def setup_status():
    status = llm_manager.get_status()
    status['progress'] = llm_manager.setup_progress
    return jsonify(status)

@app.route('/auto_setup', methods=['POST'])
@rate_limit
def auto_setup():
    try:
        data = request.get_json()
        if not data or not isinstance(data, dict):
            return jsonify({'error': 'Invalid request'}), 400
        
        model_name = data.get('model_name', 'qwen2.5:14b')
        
        if not isinstance(model_name, str):
            return jsonify({'error': 'Invalid model name type'}), 400
        
        if not validate_model_name(model_name):
            return jsonify({'error': 'Invalid model name'}), 400
        
        logger.info(f"Auto-setup requested for model: {model_name}")
        
        thread = threading.Thread(
            target=llm_manager.auto_setup,
            args=(model_name,)
        )
        thread.daemon = True
        thread.start()
        
        return jsonify({'status': 'started'})
        
    except Exception as e:
        logger.error(f"Auto-setup endpoint error: {e}")
        return jsonify({'error': str(e)[:200]}), 500

@app.route('/list_models', methods=['GET'])
def list_models():
    try:
        response = requests.get(f'{OLLAMA_URL}/api/tags', timeout=5)
        if response.status_code == 200:
            data = response.json()
            models = data.get('models', [])
            valid_models = [
                m['name'] for m in models 
                if isinstance(m.get('name'), str) and validate_model_name(m['name'])
            ]
            return jsonify({'models': valid_models})
    except (requests.RequestException, json.JSONDecodeError) as e:
        logger.error(f"Failed to list models: {e}")
    return jsonify({'models': []})

@app.route('/set_model', methods=['POST'])
@rate_limit
def set_model():
    try:
        data = request.get_json()
        if not data or not isinstance(data, dict):
            return jsonify({'error': 'Invalid request'}), 400
        
        model_name = data.get('model_name')
        
        if not isinstance(model_name, str):
            return jsonify({'error': 'Invalid model name type'}), 400
        
        if not validate_model_name(model_name):
            return jsonify({'error': 'Invalid model name'}), 400
        
        llm_manager.config['model_name'] = model_name
        llm_manager.save_config()
        logger.info(f"Model changed to: {model_name}")
        return jsonify({'success': True})
        
    except Exception as e:
        logger.error(f"Set model error: {e}")
        return jsonify({'error': str(e)[:200]}), 500

@app.route('/pull_model', methods=['POST'])
@rate_limit
def pull_model():
    """Check if model is available, download if not"""
    try:
        data = request.get_json()
        if not data or not isinstance(data, dict):
            return jsonify({'error': 'Invalid request'}), 400

        model_name = data.get('model_name')
        if not isinstance(model_name, str):
            return jsonify({'error': 'Invalid model name'}), 400

        if not validate_model_name(model_name):
            return jsonify({'error': 'Invalid model name'}), 400

        # Check if already downloading
        if llm_manager.model_download_progress.get('active'):
            return jsonify({'error': 'A download is already in progress'}), 409

        # Check if model already available
        if llm_manager.check_model_available(model_name):
            llm_manager.config['model_name'] = model_name
            llm_manager.save_config()
            logger.info(f"Model {model_name} already available, switched")
            return jsonify({'success': True, 'already_available': True})

        # Start download in background
        thread = threading.Thread(
            target=llm_manager.pull_model_async,
            args=(model_name,),
            daemon=True
        )
        thread.start()
        logger.info(f"Started downloading model: {model_name}")
        return jsonify({'success': True, 'downloading': True})

    except Exception as e:
        logger.error(f"Pull model error: {e}")
        return jsonify({'error': str(e)[:200]}), 500


@app.route('/pull_model_status', methods=['GET'])
def pull_model_status():
    """Get current model download progress"""
    return jsonify(llm_manager.model_download_progress)


@app.route('/get_prompt', methods=['GET'])
def get_prompt():
    return jsonify({'prompt': read_prompt()})

@app.route('/save_prompt', methods=['POST'])
@rate_limit
def update_prompt():
    try:
        data = request.get_json()
        if not data or not isinstance(data, dict) or 'prompt' not in data:
            return jsonify({'error': 'Invalid request'}), 400

        if not isinstance(data['prompt'], str):
            return jsonify({'error': 'Invalid prompt type'}), 400

        save_prompt(data['prompt'])
        return jsonify({'status': 'success'})

    except Exception as e:
        logger.error(f"Save prompt error: {e}")
        return jsonify({'error': str(e)[:200]}), 500

@app.route('/upload_template', methods=['POST'])
@rate_limit
def upload_template():
    """Upload a Word document template"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        original_name = file.filename
        if not original_name or not isinstance(original_name, str):
            return jsonify({'error': 'Invalid filename'}), 400

        if len(original_name) > 255:
            return jsonify({'error': 'Filename too long'}), 400

        if '..' in original_name or '/' in original_name or '\\' in original_name:
            return jsonify({'error': 'Invalid filename'}), 400

        if not original_name.lower().endswith('.docx'):
            return jsonify({'error': 'Only .docx files are allowed'}), 400

        # Save to a temp path first for validation
        temp_path = os.path.join(TEMPLATE_FOLDER, f"temp_{secrets.token_hex(8)}.docx")
        file.save(temp_path)

        # Check file size (50MB limit for templates)
        if os.path.getsize(temp_path) > 50 * 1024 * 1024:
            safe_remove_file(temp_path)
            return jsonify({'error': 'Template too large (max 50MB)'}), 400

        # Validate it's a real .docx by opening with python-docx
        try:
            from docx import Document
        except ImportError:
            safe_remove_file(temp_path)
            logger.error("python-docx is not installed")
            return jsonify({'error': 'python-docx is not installed. Run: pip install python-docx'}), 500

        try:
            Document(temp_path)
        except Exception as e:
            safe_remove_file(temp_path)
            logger.error(f"Invalid Word document: {e}")
            return jsonify({'error': f'Invalid Word document: {str(e)[:100]}'}), 400

        # Move to final location (overwrite any existing template)
        if os.path.exists(TEMPLATE_FILE):
            safe_remove_file(TEMPLATE_FILE)
        os.rename(temp_path, TEMPLATE_FILE)

        try:
            os.chmod(TEMPLATE_FILE, 0o600)
        except OSError:
            pass

        # Save original filename in config
        llm_manager.config['template_filename'] = secure_filename(original_name)
        llm_manager.save_config()

        logger.info(f"Template uploaded: {original_name}")
        return jsonify({
            'success': True,
            'filename': llm_manager.config['template_filename']
        })

    except Exception as e:
        logger.error(f"Template upload error: {e}")
        return jsonify({'error': str(e)[:200]}), 500


@app.route('/get_template_status', methods=['GET'])
def get_template_status():
    """Check if a Word template is uploaded"""
    has_template = os.path.isfile(TEMPLATE_FILE)
    filename = llm_manager.config.get('template_filename', '') if has_template else ''
    return jsonify({
        'has_template': has_template,
        'filename': filename
    })


@app.route('/delete_template', methods=['POST'])
@rate_limit
def delete_template():
    """Remove the uploaded Word template"""
    try:
        if os.path.isfile(TEMPLATE_FILE):
            safe_remove_file(TEMPLATE_FILE)

        llm_manager.config.pop('template_filename', None)
        llm_manager.save_config()

        logger.info("Template deleted")
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Template delete error: {e}")
        return jsonify({'error': str(e)[:200]}), 500


@app.route('/get_avg_processing_time', methods=['GET'])
def avg_processing_time():
    """Return the average PDF processing time"""
    avg = get_average_processing_time()
    count = len(load_processing_times())
    return jsonify({'avg_seconds': avg, 'count': count})


@app.route('/save_api_key', methods=['POST'])
@rate_limit
def save_api_key_endpoint():
    """Save API key for selected AI provider"""
    try:
        data = request.get_json()
        if not data or not isinstance(data, dict):
            return jsonify({'error': 'Invalid request'}), 400

        provider = data.get('provider')
        api_key = data.get('api_key')

        if not provider or not isinstance(provider, str):
            return jsonify({'error': 'Invalid provider'}), 400

        if not api_key or not isinstance(api_key, str):
            return jsonify({'error': 'Invalid API key'}), 400

        # Validate provider
        if provider not in ['chatgpt', 'perplexity']:
            return jsonify({'error': 'Invalid provider. Must be chatgpt or perplexity'}), 400

        # Sanitize inputs
        provider = sanitize_input(provider, 50)
        api_key = api_key.strip()[:500]  # Limit key length but don't sanitize special chars

        if save_api_key(provider, api_key):
            logger.info(f"API key saved for provider: {provider}")
            return jsonify({'success': True})
        else:
            return jsonify({'error': 'Failed to save API key'}), 500

    except Exception as e:
        logger.error(f"Save API key error: {e}")
        return jsonify({'error': str(e)[:200]}), 500

@app.route('/get_api_key', methods=['POST'])
def get_api_key_endpoint():
    """Get API key for selected AI provider"""
    try:
        data = request.get_json()
        if not data or not isinstance(data, dict):
            return jsonify({'error': 'Invalid request'}), 400

        provider = data.get('provider')

        if not provider or not isinstance(provider, str):
            return jsonify({'error': 'Invalid provider'}), 400

        provider = sanitize_input(provider, 50)
        api_key = get_api_key(provider)

        return jsonify({
            'success': True,
            'api_key': api_key if api_key else None
        })

    except Exception as e:
        logger.error(f"Get API key error: {e}")
        return jsonify({'error': str(e)[:200]}), 500

@app.route('/set_provider', methods=['POST'])
@rate_limit
def set_provider_endpoint():
    """Set the active LLM provider"""
    try:
        data = request.get_json()
        if not data or not isinstance(data, dict):
            return jsonify({'error': 'Invalid request'}), 400

        provider = data.get('provider')
        api_model = data.get('api_model')

        if not provider or not isinstance(provider, str):
            return jsonify({'error': 'Invalid provider'}), 400

        # Validate provider
        if provider not in ['ollama', 'chatgpt', 'perplexity']:
            return jsonify({'error': 'Invalid provider'}), 400

        provider = sanitize_input(provider, 50)

        llm_manager.config['provider'] = provider

        if api_model and isinstance(api_model, str):
            llm_manager.config['api_model'] = sanitize_input(api_model, 100)

        llm_manager.save_config()
        logger.info(f"Provider changed to: {provider}")

        return jsonify({'success': True})

    except Exception as e:
        logger.error(f"Set provider error: {e}")
        return jsonify({'error': str(e)[:200]}), 500

@app.route('/get_provider', methods=['GET'])
def get_provider_endpoint():
    """Get the current LLM provider and model"""
    try:
        return jsonify({
            'success': True,
            'provider': llm_manager.config.get('provider', 'ollama'),
            'api_model': llm_manager.config.get('api_model', 'gpt-4o-mini'),
            'model_name': llm_manager.config.get('model_name', 'qwen2.5:14b')
        })
    except Exception as e:
        logger.error(f"Get provider error: {e}")
        return jsonify({'error': str(e)[:200]}), 500

@app.route('/test_api_connection', methods=['POST'])
@rate_limit
def test_api_connection():
    """Test connection to selected AI provider using API key"""
    try:
        data = request.get_json()
        if not data or not isinstance(data, dict):
            return jsonify({'error': 'Invalid request'}), 400

        provider = data.get('provider')
        model = data.get('model', 'gpt-4o-mini')

        if not provider or not isinstance(provider, str):
            return jsonify({'error': 'Invalid provider'}), 400

        provider = sanitize_input(provider, 50)
        model = sanitize_input(model, 100)

        api_key = get_api_key(provider)

        if not api_key:
            return jsonify({
                'success': False,
                'error': 'No API key found. Please save your API key first.'
            })

        # Test connection based on provider
        if provider == 'chatgpt':
            headers = {
                'Authorization': f'Bearer {api_key}',
                'Content-Type': 'application/json'
            }

            test_payload = {
                'model': model,
                'messages': [
                    {'role': 'user', 'content': 'Hello'}
                ],
                'max_tokens': 5
            }

            response = requests.post(
                'https://api.openai.com/v1/chat/completions',
                headers=headers,
                json=test_payload,
                timeout=10
            )

            if response.status_code == 200:
                return jsonify({
                    'success': True,
                    'provider': provider,
                    'model': model
                })
            else:
                error_detail = response.text[:200]
                return jsonify({
                    'success': False,
                    'error': f'API error: {response.status_code} - {error_detail}'
                })

        elif provider == 'perplexity':
            headers = {
                'Authorization': f'Bearer {api_key}',
                'Content-Type': 'application/json'
            }

            test_payload = {
                'model': model,
                'messages': [
                    {'role': 'system', 'content': 'Be brief.'},
                    {'role': 'user', 'content': 'Hi'}
                ]
            }

            response = requests.post(
                'https://api.perplexity.ai/chat/completions',
                headers=headers,
                json=test_payload,
                timeout=10
            )

            if response.status_code == 200:
                return jsonify({
                    'success': True,
                    'provider': provider,
                    'model': model
                })
            else:
                error_detail = response.text[:200]
                return jsonify({
                    'success': False,
                    'error': f'API error: {response.status_code} - {error_detail}'
                })

        else:
            return jsonify({
                'success': False,
                'error': 'Invalid provider'
            })

    except requests.exceptions.RequestException as e:
        logger.error(f"Test API connection error: {e}")
        return jsonify({
            'success': False,
            'error': f'Connection error: {str(e)[:100]}'
        })
    except Exception as e:
        logger.error(f"Test API connection error: {e}")
        return jsonify({'error': str(e)[:200]}), 500

@app.route('/process_pdf', methods=['POST'])
@rate_limit
def process_pdf():
    filepath = None
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not validate_filename(file.filename):
            return jsonify({'error': 'Invalid filename'}), 400
        
        filename = secure_filename(file.filename)
        file_ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
        if file_ext not in ALLOWED_EXTENSIONS:
            return jsonify({'error': 'Only PDF and ZIP files are allowed'}), 400

        unique_filename = f"{secrets.token_hex(8)}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)

        logger.info(f"Processing file: {filename}")
        file.save(filepath)

        if os.path.getsize(filepath) > app.config['MAX_CONTENT_LENGTH']:
            safe_remove_file(filepath)
            return jsonify({'error': 'File too large (max 50MB)'}), 400

        try:
            start_time = time.time()

            if file_ext == 'zip':
                document_text = extract_latex_text(filepath)
            else:
                document_text = extract_pdf_text(filepath)
            prompt = read_prompt()

            # Augment prompt with template structure if a template is uploaded
            has_template = os.path.isfile(TEMPLATE_FILE)
            if has_template:
                template_structure = extract_template_structure(TEMPLATE_FILE)
                if template_structure:
                    prompt = f"{prompt}\n\n{template_structure}"

            llm_response = query_llm(prompt, document_text)

            if has_template:
                # Generate .docx output using the template
                output_filename = f"{Path(filename).stem}_output.docx"
                output_path = os.path.join(
                    app.config['OUTPUT_FOLDER'],
                    f"{secrets.token_hex(8)}_{output_filename}"
                )
                doc = generate_docx_from_template(TEMPLATE_FILE, llm_response)
                doc.save(output_path)
            else:
                # Plain text output (no template)
                output_filename = f"{Path(filename).stem}_output.txt"
                output_path = os.path.join(
                    app.config['OUTPUT_FOLDER'],
                    f"{secrets.token_hex(8)}_{output_filename}"
                )
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(llm_response)

            duration = round(time.time() - start_time, 2)
            save_processing_time(duration)

            try:
                os.chmod(output_path, 0o600)
            except OSError:
                pass

            output_id = os.path.basename(output_path)
            logger.info(f"PDF processed successfully: {output_id} ({duration}s)")

            return jsonify({
                'status': 'success',
                'output_file': output_id,
                'original_name': output_filename,
                'duration': duration
            })
        
        finally:
            if filepath:
                safe_remove_file(filepath)
                
    except Exception as e:
        logger.error(f"Process PDF error: {e}")
        if filepath:
            safe_remove_file(filepath)
        return jsonify({'error': str(e)[:200]}), 500

@app.route('/download/<filename>')
def download(filename):
    try:
        safe_filename = secure_filename(filename)
        if not safe_filename or safe_filename != filename:
            return jsonify({'error': 'Invalid filename'}), 400
        
        filepath = os.path.join(app.config['OUTPUT_FOLDER'], safe_filename)
        
        if not os.path.isfile(filepath):
            return jsonify({'error': 'File not found'}), 404
        
        real_path = os.path.abspath(filepath)
        output_dir = os.path.abspath(app.config['OUTPUT_FOLDER'])
        
        if not real_path.startswith(output_dir):
            logger.warning(f"Path traversal attempt: {filename}")
            return jsonify({'error': 'Access denied'}), 403
        
        logger.info(f"Downloading file: {safe_filename}")
        response = send_file(filepath, as_attachment=True)
        
        @response.call_on_close
        def cleanup():
            safe_remove_file(filepath)
        
        return response
        
    except Exception as e:
        logger.error(f"Download error: {e}")
        return jsonify({'error': str(e)[:200]}), 500

def cleanup_old_files():
    """Remove files older than 1 hour"""
    try:
        current_time = time.time()
        for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
            if not os.path.exists(folder):
                continue
            
            cleaned = 0
            for filename in os.listdir(folder):
                filepath = os.path.join(folder, filename)
                if os.path.isfile(filepath):
                    if current_time - os.path.getmtime(filepath) > 3600:
                        if safe_remove_file(filepath):
                            cleaned += 1
            
            if cleaned > 0:
                logger.info(f"Cleaned {cleaned} old files from {folder}")
                
    except OSError as e:
        logger.error(f"Cleanup error: {e}")

def signal_handler(sig, frame):
    """Handle shutdown signals"""
    logger.info("Shutting down gracefully...")
    print("\nShutting down gracefully...")
    llm_manager.cleanup()
    cleanup_old_files()
    sys.exit(0)

def startup_sequence():
    """Startup sequence"""
    print("=== LLM PDF Processor Startup ===")
    print(f"Python: {sys.version}")
    print(f"Working directory: {os.getcwd()}")
    logger.info("Application starting")
    
    cleanup_old_files()
    
    if llm_manager.config.get('setup_complete', False):
        print("✓ Setup complete")
        logger.info("Setup already complete")
        
        if llm_manager.start_ollama_service():
            print("✓ Ollama service started")
            print(f"✓ Model '{llm_manager.config['model_name']}' ready")
            print("\n=== System Ready ===")
            print("Access the application at:")
            print("  http://127.0.0.1:5000")
            print("  http://localhost:5000")
            if 'WSL' in platform.uname().release or 'microsoft' in platform.uname().release.lower():
                print("\nWSL Detected: Access from Windows browser at:")
                print("  http://localhost:5000")
            logger.info("System ready")
        else:
            print("✗ Failed to start Ollama service")
            print("Complete setup at http://127.0.0.1:5000")
            logger.warning("Ollama service failed to start")
    else:
        print("First-time setup required")
        print("Access the application at http://127.0.0.1:5000")
        logger.info("First-time setup required")
    
    print("\nPress Ctrl+C to stop")

if __name__ == '__main__':
    signal.signal(signal.SIGINT, signal_handler)
    signal.signal(signal.SIGTERM, signal_handler)
    
    startup_sequence()
    
    try:
        logger.info("Starting Flask application on 0.0.0.0:5000")
        app.run(debug=False, host='0.0.0.0', port=5000, use_reloader=False, threaded=True)
    except KeyboardInterrupt:
        signal_handler(None, None)
    except Exception as e:
        logger.error(f"Application error: {e}")
        raise
